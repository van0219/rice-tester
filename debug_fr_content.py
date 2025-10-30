#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# Debug FR content removal
template_path = r"d:\AmazonQ\InforQ\RICE_Tester\TES-070-Template\TES-070_Custom_Extension_Unit_Test_Results_v2.0.docx"
doc = Document(template_path)

scenarios = [("dummy",)]  # 1 scenario
fr_sections_to_remove = ['FR 1.2', 'FR 1.3']

print("=== DEBUGGING FR CONTENT REMOVAL ===")
print(f"Total paragraphs: {len(doc.paragraphs)}")

paragraphs_to_remove = []
i = 0
while i < len(doc.paragraphs):
    para = doc.paragraphs[i]
    
    # Check if this FR section should be removed
    for fr_to_remove in fr_sections_to_remove:
        if para.text.startswith(fr_to_remove):
            print(f"\nFound FR section to remove at line {i}: '{para.text}'")
            
            # Mark this FR section and all its content for removal
            j = i
            content_lines = []
            while j < len(doc.paragraphs):
                content_lines.append(f"  Line {j}: '{doc.paragraphs[j].text[:60]}...'")
                
                # Stop when we hit the next FR section or section 4
                if j > i and (doc.paragraphs[j].text.startswith('FR ') or doc.paragraphs[j].text.startswith('4\t')):
                    print(f"  Stopping at line {j}: '{doc.paragraphs[j].text[:60]}...'")
                    break
                paragraphs_to_remove.append(j)
                j += 1
            
            print(f"  Content to remove ({len(content_lines)} lines):")
            for line in content_lines[:10]:  # Show first 10 lines
                print(line)
            if len(content_lines) > 10:
                print(f"    ... and {len(content_lines) - 10} more lines")
            
            i = j - 1  # Skip to after this section
            break
    i += 1

print(f"\nTotal paragraphs to remove: {len(paragraphs_to_remove)}")
print(f"Unique paragraphs to remove: {len(set(paragraphs_to_remove))}")

# Show what would be removed
if paragraphs_to_remove:
    print("\nFirst 5 paragraphs that would be removed:")
    for idx in sorted(set(paragraphs_to_remove))[:5]:
        if idx < len(doc.paragraphs):
            print(f"  Line {idx}: '{doc.paragraphs[idx].text[:60]}...'")