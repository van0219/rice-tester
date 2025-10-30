#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
os.system('chcp 65001 >nul 2>&1')  # Set Windows console to UTF-8
sys.path.append(os.path.dirname(__file__))

from database_manager import DatabaseManager
from docx import Document
from io import BytesIO

def debug_tes070_structure():
    """Debug the structure of the latest TES-070 document"""
    
    print("Debugging TES-070 document structure...")
    
    try:
        # Initialize database manager
        db_manager = DatabaseManager(user_id=1)
        
        # Get latest TES-070 version for INT001
        versions = db_manager.get_tes070_versions("INT001")
        
        if not versions:
            print("No TES-070 versions found")
            return False
        
        latest_version = versions[0]
        version_id = latest_version[0]
        
        # Get document content
        file_content = db_manager.get_tes070_content(version_id)
        
        # Load document from bytes
        doc_stream = BytesIO(file_content)
        doc = Document(doc_stream)
        
        print("Document structure analysis:")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        
        # Check first 10 paragraphs
        print("\nFirst 10 paragraphs:")
        for i, paragraph in enumerate(doc.paragraphs[:10]):
            text = paragraph.text.strip()
            has_images = bool(paragraph._element.xpath('.//a:blip'))
            
            if text or has_images:
                print(f"  Paragraph {i+1}: {'[IMAGE]' if has_images else ''} '{text[:50]}{'...' if len(text) > 50 else ''}'")
            else:
                print(f"  Paragraph {i+1}: [EMPTY]")
        
        # Check if logo files exist
        logo_ico = os.path.join(os.path.dirname(__file__), 'infor_logo.ico')
        logo_png = os.path.join(os.path.dirname(__file__), 'infor_logo.png')
        
        print(f"\nLogo files status:")
        print(f"  infor_logo.ico: {'EXISTS' if os.path.exists(logo_ico) else 'MISSING'}")
        print(f"  infor_logo.png: {'EXISTS' if os.path.exists(logo_png) else 'MISSING'}")
        
        if os.path.exists(logo_png):
            print(f"  PNG size: {os.path.getsize(logo_png)} bytes")
        if os.path.exists(logo_ico):
            print(f"  ICO size: {os.path.getsize(logo_ico)} bytes")
            
        return True
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    debug_tes070_structure()