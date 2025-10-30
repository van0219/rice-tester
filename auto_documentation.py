#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import json
import tkinter as tk
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class AutoDocumentation:
    """Automatic documentation generation for RICE Tester scenarios"""
    
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.reports_dir = os.path.join(os.path.dirname(__file__), 'reports')
        self._ensure_reports_directory()
    
    def _ensure_reports_directory(self):
        """Create reports directory if it doesn't exist"""
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)
    
    def generate_scenario_documentation(self, scenario_id, rice_profile):
        """Generate comprehensive documentation for a scenario"""
        try:
            # Get scenario data
            scenario_data = self._get_scenario_data(scenario_id, rice_profile)
            if not scenario_data:
                return None
            
            # Generate DOCX report
            doc_path = self._generate_docx_report(scenario_data)
            
            # Generate Markdown report
            md_path = self._generate_markdown_report(scenario_data)
            
            # Generate JSON summary
            json_path = self._generate_json_summary(scenario_data)
            
            return {
                'docx': doc_path,
                'markdown': md_path,
                'json': json_path,
                'scenario': scenario_data['scenario']['description']
            }
            
        except Exception as e:
            print(f"Documentation generation failed: {e}")
            return None
    
    def generate_bulk_documentation(self, rice_profile):
        """Generate documentation for all scenarios in a RICE profile"""
        try:
            scenarios = self.db_manager.get_scenarios(rice_profile)
            generated_docs = []
            
            for scenario in scenarios:
                scenario_id = scenario[0]
                doc_result = self.generate_scenario_documentation(scenario_id, rice_profile)
                if doc_result:
                    generated_docs.append(doc_result)
            
            # Generate summary report
            summary_path = self._generate_bulk_summary(generated_docs, rice_profile)
            
            return {
                'individual_docs': generated_docs,
                'summary_report': summary_path,
                'total_scenarios': len(generated_docs)
            }
            
        except Exception as e:
            print(f"Bulk documentation generation failed: {e}")
            return None
    
    def _get_scenario_data(self, scenario_id, rice_profile):
        """Get comprehensive scenario data from database"""
        try:
            cursor = self.db_manager.conn.cursor()
            
            # Get scenario info
            cursor.execute("""
                SELECT scenario_number, description, result, executed_at, created_at
                FROM scenarios 
                WHERE id = ? AND user_id = ?
            """, (scenario_id, self.db_manager.user_id))
            
            scenario_info = cursor.fetchone()
            if not scenario_info:
                return None
            
            # Get scenario steps
            cursor.execute("""
                SELECT ss.step_order, 
                       COALESCE(ts.name, ss.step_name) as step_name,
                       COALESCE(ts.step_type, ss.step_type) as step_type,
                       CASE 
                           WHEN COALESCE(ts.step_type, ss.step_type) IN ('Wait', 'Text Input') 
                           THEN COALESCE(NULLIF(ss.step_description, ''), NULLIF(ss.custom_value, 'None'), ts.default_value)
                           ELSE COALESCE(ts.target, ss.step_target)
                       END as step_target,
                       ss.execution_status,
                       ss.screenshot_timestamp
                FROM scenario_steps ss
                LEFT JOIN test_steps ts ON ss.test_step_id = ts.id
                WHERE ss.user_id = ? AND ss.rice_profile = ? AND ss.scenario_number = ?
                ORDER BY ss.step_order
            """, (self.db_manager.user_id, str(rice_profile), scenario_info[0]))
            
            steps = cursor.fetchall()
            
            # Get RICE profile info
            cursor.execute("""
                SELECT profile_name, base_url, description
                FROM rice_profiles 
                WHERE id = ? AND user_id = ?
            """, (rice_profile, self.db_manager.user_id))
            
            profile_info = cursor.fetchone()
            
            return {
                'scenario': {
                    'id': scenario_id,
                    'number': scenario_info[0],
                    'description': scenario_info[1],
                    'result': scenario_info[2],
                    'executed_at': scenario_info[3],
                    'created_at': scenario_info[4]
                },
                'steps': [
                    {
                        'order': step[0],
                        'name': step[1],
                        'type': step[2],
                        'target': step[3],
                        'status': step[4],
                        'screenshot_timestamp': step[5]
                    } for step in steps
                ],
                'profile': {
                    'id': rice_profile,
                    'name': profile_info[0] if profile_info else 'Unknown',
                    'base_url': profile_info[1] if profile_info else '',
                    'description': profile_info[2] if profile_info else ''
                }
            }
            
        except Exception as e:
            print(f"Error getting scenario data: {e}")
            return None
    
    def _generate_docx_report(self, scenario_data):
        """Generate DOCX report for scenario"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(f"Scenario #{scenario_data['scenario']['number']} - Test Documentation", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Scenario Information
            doc.add_heading('Scenario Information', level=1)
            
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Scenario Number', str(scenario_data['scenario']['number'])),
                ('Description', scenario_data['scenario']['description']),
                ('RICE Profile', scenario_data['profile']['name']),
                ('Base URL', scenario_data['profile']['base_url']),
                ('Status', scenario_data['scenario']['result'] or 'Not Run'),
                ('Last Executed', scenario_data['scenario']['executed_at'] or 'Never')
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = str(value)
            
            # Test Steps
            doc.add_heading('Test Steps', level=1)
            
            if scenario_data['steps']:
                steps_table = doc.add_table(rows=len(scenario_data['steps']) + 1, cols=5)
                steps_table.style = 'Table Grid'
                
                # Headers
                headers = ['Step', 'Name', 'Type', 'Target/Value', 'Status']
                for i, header in enumerate(headers):
                    steps_table.cell(0, i).text = header
                
                # Step data
                for i, step in enumerate(scenario_data['steps'], 1):
                    steps_table.cell(i, 0).text = str(step['order'])
                    steps_table.cell(i, 1).text = step['name'] or ''
                    steps_table.cell(i, 2).text = step['type'] or ''
                    steps_table.cell(i, 3).text = step['target'] or ''
                    steps_table.cell(i, 4).text = step['status'] or 'Pending'
            else:
                doc.add_paragraph('No steps defined for this scenario.')
            
            # Execution Summary
            doc.add_heading('Execution Summary', level=1)
            
            total_steps = len(scenario_data['steps'])
            completed_steps = len([s for s in scenario_data['steps'] if s['status'] == 'completed'])
            
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Total Steps: {total_steps}\n")
            summary_para.add_run(f"Completed Steps: {completed_steps}\n")
            summary_para.add_run(f"Success Rate: {(completed_steps/total_steps*100):.1f}%\n" if total_steps > 0 else "Success Rate: N/A\n")
            summary_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Save document
            filename = f"Scenario_{scenario_data['scenario']['number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc_path = os.path.join(self.reports_dir, filename)
            doc.save(doc_path)
            
            return doc_path
            
        except Exception as e:
            print(f"DOCX generation failed: {e}")
            return None
    
    def _generate_markdown_report(self, scenario_data):
        """Generate Markdown report for scenario"""
        try:
            md_content = []
            
            # Title
            md_content.append(f"# Scenario #{scenario_data['scenario']['number']} - Test Documentation\n")
            
            # Scenario Information
            md_content.append("## Scenario Information\n")
            md_content.append(f"- **Scenario Number:** {scenario_data['scenario']['number']}")
            md_content.append(f"- **Description:** {scenario_data['scenario']['description']}")
            md_content.append(f"- **RICE Profile:** {scenario_data['profile']['name']}")
            md_content.append(f"- **Base URL:** {scenario_data['profile']['base_url']}")
            md_content.append(f"- **Status:** {scenario_data['scenario']['result'] or 'Not Run'}")
            md_content.append(f"- **Last Executed:** {scenario_data['scenario']['executed_at'] or 'Never'}\n")
            
            # Test Steps
            md_content.append("## Test Steps\n")
            
            if scenario_data['steps']:
                md_content.append("| Step | Name | Type | Target/Value | Status |")
                md_content.append("|------|------|------|--------------|--------|")
                
                for step in scenario_data['steps']:
                    md_content.append(f"| {step['order']} | {step['name'] or ''} | {step['type'] or ''} | {step['target'] or ''} | {step['status'] or 'Pending'} |")
            else:
                md_content.append("No steps defined for this scenario.")
            
            md_content.append("")
            
            # Execution Summary
            md_content.append("## Execution Summary\n")
            
            total_steps = len(scenario_data['steps'])
            completed_steps = len([s for s in scenario_data['steps'] if s['status'] == 'completed'])
            
            md_content.append(f"- **Total Steps:** {total_steps}")
            md_content.append(f"- **Completed Steps:** {completed_steps}")
            md_content.append(f"- **Success Rate:** {(completed_steps/total_steps*100):.1f}%" if total_steps > 0 else "- **Success Rate:** N/A")
            md_content.append(f"- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Save markdown
            filename = f"Scenario_{scenario_data['scenario']['number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            md_path = os.path.join(self.reports_dir, filename)
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_content))
            
            return md_path
            
        except Exception as e:
            print(f"Markdown generation failed: {e}")
            return None
    
    def _generate_json_summary(self, scenario_data):
        """Generate JSON summary for scenario"""
        try:
            summary = {
                'scenario': scenario_data['scenario'],
                'profile': scenario_data['profile'],
                'steps': scenario_data['steps'],
                'summary': {
                    'total_steps': len(scenario_data['steps']),
                    'completed_steps': len([s for s in scenario_data['steps'] if s['status'] == 'completed']),
                    'success_rate': (len([s for s in scenario_data['steps'] if s['status'] == 'completed']) / len(scenario_data['steps']) * 100) if scenario_data['steps'] else 0,
                    'generated_at': datetime.now().isoformat()
                }
            }
            
            # Save JSON
            filename = f"Scenario_{scenario_data['scenario']['number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            json_path = os.path.join(self.reports_dir, filename)
            
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(summary, f, indent=2, ensure_ascii=False)
            
            return json_path
            
        except Exception as e:
            print(f"JSON generation failed: {e}")
            return None
    
    def _generate_bulk_summary(self, generated_docs, rice_profile):
        """Generate summary report for bulk documentation"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(f'RICE Profile Documentation Summary', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary Information
            doc.add_heading('Summary Information', level=1)
            
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ('RICE Profile ID', str(rice_profile)),
                ('Total Scenarios Documented', str(len(generated_docs))),
                ('Generated Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Generated By', 'RICE Tester Auto-Documentation')
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = label
                summary_table.cell(i, 1).text = value
            
            # Scenarios List
            doc.add_heading('Documented Scenarios', level=1)
            
            if generated_docs:
                scenarios_table = doc.add_table(rows=len(generated_docs) + 1, cols=3)
                scenarios_table.style = 'Table Grid'
                
                # Headers
                scenarios_table.cell(0, 0).text = 'Scenario'
                scenarios_table.cell(0, 1).text = 'Description'
                scenarios_table.cell(0, 2).text = 'Files Generated'
                
                # Scenario data
                for i, doc_info in enumerate(generated_docs, 1):
                    scenarios_table.cell(i, 0).text = f"Scenario #{i}"
                    scenarios_table.cell(i, 1).text = doc_info['scenario']
                    scenarios_table.cell(i, 2).text = "DOCX, Markdown, JSON"
            
            # Save summary
            filename = f"RICE_Profile_{rice_profile}_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            summary_path = os.path.join(self.reports_dir, filename)
            doc.save(summary_path)
            
            return summary_path
            
        except Exception as e:
            print(f"Bulk summary generation failed: {e}")
            return None
    
    def get_reports_directory(self):
        """Get the reports directory path"""
        return self.reports_dir
    
    def cleanup_old_reports(self, days_old=30):
        """Clean up reports older than specified days"""
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (days_old * 24 * 60 * 60)
            
            cleaned_count = 0
            for filename in os.listdir(self.reports_dir):
                file_path = os.path.join(self.reports_dir, filename)
                if os.path.isfile(file_path):
                    file_time = os.path.getmtime(file_path)
                    if file_time < cutoff_time:
                        os.remove(file_path)
                        cleaned_count += 1
            
            return cleaned_count
            
        except Exception as e:
            print(f"Cleanup failed: {e}")
            return 0

# Integration helper for scenario forms
def add_documentation_button_to_scenario(scenario_frame, scenario_id, rice_profile, db_manager, show_popup_callback):
    """Add documentation generation button to scenario interface"""
    
    def generate_docs():
        try:
            doc_generator = AutoDocumentation(db_manager)
            result = doc_generator.generate_scenario_documentation(scenario_id, rice_profile)
            
            if result:
                show_popup_callback("Documentation Generated", 
                                  f"Documentation created successfully!\n\nFiles generated:\n• DOCX Report\n• Markdown Summary\n• JSON Data\n\nSaved to: {doc_generator.get_reports_directory()}", 
                                  "success")
            else:
                show_popup_callback("Generation Failed", "Failed to generate documentation. Please try again.", "error")
                
        except Exception as e:
            show_popup_callback("Error", f"Documentation generation error: {str(e)}", "error")
    
    doc_btn = tk.Button(scenario_frame, text="📄 Generate Docs", font=('Segoe UI', 9, 'bold'),
                       bg='#8b5cf6', fg='#ffffff', relief='flat', padx=10, pady=4,
                       cursor='hand2', bd=0, command=generate_docs)
    
    return doc_btn
