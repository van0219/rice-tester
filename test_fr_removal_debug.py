#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# Test FR removal logic
template_path = r"d:\AmazonQ\InforQ\RICE_Tester\TES-070-Template\TES-070_Custom_Extension_Unit_Test_Results_v2.0.docx"
doc = Document(template_path)

print("=== BEFORE FR REMOVAL ===")
fr_count = 0
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith('FR 1.'):
        print(f"Line {i}: {para.text}")
        fr_count += 1

print(f"Total FR sections found: {fr_count}")

# Simulate 1 scenario (should remove FR 1.2 and FR 1.3)
scenarios = [("dummy",)]  # 1 scenario
print(f"\nScenarios count: {len(scenarios)}")

if len(scenarios) < 3:
    fr_sections_to_remove = []
    if len(scenarios) == 1:
        fr_sections_to_remove = ['FR 1.2', 'FR 1.3']
    elif len(scenarios) == 2:
        fr_sections_to_remove = ['FR 1.3']
    
    print(f"FR sections to remove: {fr_sections_to_remove}")
    
    # Find and remove unused FR sections from content
    paragraphs_to_remove = []
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        
        # Check if this FR section should be removed
        for fr_to_remove in fr_sections_to_remove:
            if para.text.startswith(fr_to_remove) and (not hasattr(para, 'style') or not para.style or para.style.name != 'toc 2'):
                print(f"Found FR section to remove: {para.text}")
                # Mark this FR section and all its content for removal
                j = i
                while j < len(doc.paragraphs):
                    # Stop when we hit the next FR section or section 4
                    if j > i and (doc.paragraphs[j].text.startswith('FR ') or doc.paragraphs[j].text.startswith('4\t')):
                        break
                    paragraphs_to_remove.append(j)
                    j += 1
                i = j - 1  # Skip to after this section
                break
        i += 1
    
    print(f"Paragraphs to remove: {len(paragraphs_to_remove)} paragraphs")
    
    # Remove content paragraphs in reverse order to maintain indices
    for para_index in reversed(sorted(set(paragraphs_to_remove))):
        if para_index < len(doc.paragraphs):
            p = doc.paragraphs[para_index]
            print(f"Removing paragraph {para_index}: {p.text[:50]}...")
            p._element.getparent().remove(p._element)

print("\n=== AFTER FR REMOVAL ===")
fr_count_after = 0
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith('FR 1.'):
        print(f"Line {i}: {para.text}")
        fr_count_after += 1

print(f"Total FR sections after removal: {fr_count_after}")