#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
os.system('chcp 65001 >nul 2>&1')  # Set Windows console to UTF-8

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    Image = None

def convert_logo_to_compatible_format():
    """Convert logo to a format compatible with python-docx"""
    
    print("Converting logo to compatible format...")
    
    # Check if PIL is available
    try:
        from PIL import Image
        pil_available = True
    except ImportError:
        pil_available = False
    
    if not pil_available:
        print("PIL (Pillow) not available. Installing...")
        try:
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])
            from PIL import Image
            print("Pillow installed successfully!")
            pil_available = True
        except Exception as e:
            print(f"Failed to install Pillow: {e}")
            return False
    
    # Try to convert the existing PNG
    logo_png = 'infor_logo.png'
    logo_ico = 'infor_logo.ico'
    
    if os.path.exists(logo_png):
        try:
            print(f"Converting {logo_png}...")
            
            # Open and convert PNG
            with Image.open(logo_png) as img:
                print(f"Original format: {img.format}, size: {img.size}, mode: {img.mode}")
                
                # Convert to RGB if needed (remove alpha channel)
                if img.mode in ('RGBA', 'LA', 'P'):
                    print("Converting to RGB...")
                    # Create white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Resize to reasonable size (1 inch at 96 DPI = 96 pixels)
                target_size = (96, 96)
                img = img.resize(target_size, Image.Resampling.LANCZOS)
                
                # Save as standard PNG
                new_logo = 'infor_logo_fixed.png'
                img.save(new_logo, 'PNG', optimize=True)
                print(f"Saved as {new_logo} ({os.path.getsize(new_logo)} bytes)")
                
                return new_logo
                
        except Exception as e:
            print(f"Error converting PNG: {e}")
    
    # Try ICO if PNG failed
    if os.path.exists(logo_ico):
        try:
            print(f"Converting {logo_ico}...")
            
            with Image.open(logo_ico) as img:
                print(f"ICO format: {img.format}, size: {img.size}, mode: {img.mode}")
                
                # Convert to RGB
                if img.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                
                # Resize to standard size
                target_size = (96, 96)
                img = img.resize(target_size, Image.Resampling.LANCZOS)
                
                # Save as PNG
                new_logo = 'infor_logo_from_ico.png'
                img.save(new_logo, 'PNG', optimize=True)
                print(f"Saved as {new_logo} ({os.path.getsize(new_logo)} bytes)")
                
                return new_logo
                
        except Exception as e:
            print(f"Error converting ICO: {e}")
    
    print("No logo files could be converted")
    return False

def test_converted_logo(logo_file):
    """Test if the converted logo works with python-docx"""
    
    print(f"Testing converted logo: {logo_file}")
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # Try to add the converted logo
        run.add_picture(logo_file, width=Inches(1), height=Inches(1))
        
        # Save test document
        test_file = 'test_converted_logo.docx'
        doc.save(test_file)
        
        # Verify
        test_doc = Document(test_file)
        image_count = 0
        for paragraph in test_doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//a:blip'):
                    image_count += 1
        
        print(f"Test result: {image_count} images found")
        return image_count > 0
        
    except Exception as e:
        print(f"Test failed: {e}")
        return False

if __name__ == "__main__":
    converted_logo = convert_logo_to_compatible_format()
    
    if converted_logo:
        success = test_converted_logo(converted_logo)
        if success:
            print(f"\nSUCCESS: {converted_logo} works with python-docx!")
            print("You can now use this file in the TES-070 generator.")
        else:
            print(f"\nFAILED: {converted_logo} still doesn't work with python-docx")
    else:
        print("\nFAILED: Could not convert any logo files")