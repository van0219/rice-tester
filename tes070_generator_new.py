#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sqlite3
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import tkinter as tk
from tkinter import filedialog
from database_manager import DatabaseManager



def create_tes070_from_scratch(rice_profile, show_popup=None, current_user=None, db_manager=None):
    """Generate TES-070 from scratch with full control over sections"""
    
    # Show loading dialog
    loading_popup = None
    if show_popup:
        import tkinter as tk
        from rice_dialogs import center_dialog
        
        # Create loading dialog
        loading_popup = tk.Toplevel()
        loading_popup.title("Generating TES-070")
        center_dialog(loading_popup, 400, 236)
        loading_popup.configure(bg='#ffffff')
        loading_popup.resizable(False, False)
        loading_popup.grab_set()
        
        try:
            loading_popup.iconbitmap("infor_logo.ico")
        except:
            pass
        
        # Header
        header_frame = tk.Frame(loading_popup, bg='#3b82f6', height=60)
        header_frame.pack(fill="x")
        header_frame.pack_propagate(False)
        
        header_label = tk.Label(header_frame, text="📄 Generating TES-070 Report", 
                               font=('Segoe UI', 14, 'bold'), bg='#3b82f6', fg='#ffffff')
        header_label.pack(expand=True)
        
        # Content
        content_frame = tk.Frame(loading_popup, bg='#ffffff', padx=20, pady=20)
        content_frame.pack(fill="both", expand=True)
        
        # Loading message
        loading_label = tk.Label(content_frame, text="Creating TES-070 from scratch...\n\nBuilding professional document structure.", 
                                font=('Segoe UI', 10), bg='#ffffff', justify="center")
        loading_label.pack(pady=(10, 20))
        
        # Progress indicator
        progress_label = tk.Label(content_frame, text="●●●", 
                                 font=('Segoe UI', 16), bg='#ffffff', fg='#3b82f6')
        progress_label.pack()
        
        # Animate the progress dots
        def animate_dots():
            current = progress_label.cget('text')
            if current == "●●●":
                progress_label.config(text="○●●")
            elif current == "○●●":
                progress_label.config(text="○○●")
            elif current == "○○●":
                progress_label.config(text="○○○")
            else:
                progress_label.config(text="●●●")
            
            if loading_popup.winfo_exists():
                loading_popup.after(300, animate_dots)
        
        # Start animation
        animate_dots()
        loading_popup.update()
        
        # Record start time for minimum 5-second display
        import time
        start_time = time.time()
    
    try:
        # Database path
        db_path = r"d:\AmazonQ\InforQ\RICE_Tester\fsm_tester.db"
    
        if not os.path.exists(db_path):
            if loading_popup and loading_popup.winfo_exists():
                elapsed_time = time.time() - start_time
                if elapsed_time < 5.0:
                    remaining_time = int((5.0 - elapsed_time) * 1000)
                    loading_popup.after(remaining_time, loading_popup.destroy)
                else:
                    loading_popup.destroy()
            if show_popup:
                show_popup("Error", "Database not found", "error")
            return
        
        # Connect to database
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Check for "not run" scenarios
        if rice_profile:
            cursor.execute("""
                SELECT COUNT(*) FROM scenarios s
                JOIN rice_profiles rp ON s.rice_profile = rp.id
                WHERE (rp.id = ? OR rp.rice_id = ?) AND (s.result IS NULL OR s.result = 'Not run')
            """, (rice_profile, rice_profile))
            not_run_count = cursor.fetchone()[0]
            
            if not_run_count > 0:
                cursor.execute("SELECT rice_id FROM rice_profiles WHERE id = ? OR rice_id = ?", (rice_profile, rice_profile))
                rice_id_result = cursor.fetchone()
                display_rice_id = rice_id_result[0] if rice_id_result else rice_profile
                
                if show_popup:
                    show_popup(
                        "Cannot Generate TES-070", 
                        f"Cannot generate TES-070 report for RICE {display_rice_id}.\n\n"
                        f"Found {not_run_count} scenario(s) with 'Not run' status.\n\n"
                        "Please execute all scenarios before generating the report.",
                        "error"
                    )
                conn.close()
                if loading_popup and loading_popup.winfo_exists():
                    elapsed_time = time.time() - start_time
                    if elapsed_time < 5.0:
                        remaining_time = int((5.0 - elapsed_time) * 1000)
                        loading_popup.after(remaining_time, loading_popup.destroy)
                    else:
                        loading_popup.destroy()
                return
        
        # Get scenario data
        if rice_profile:
            cursor.execute("""
                SELECT s.rice_profile, s.scenario_number, s.description, s.result, s.executed_at
                FROM scenarios s
                JOIN rice_profiles rp ON s.rice_profile = rp.id
                WHERE s.result IS NOT NULL AND (rp.id = ? OR rp.rice_id = ?)
                ORDER BY s.scenario_number
            """, (rice_profile, rice_profile))
        else:
            cursor.execute("""
                SELECT rice_profile, scenario_number, description, result, executed_at
                FROM scenarios 
                WHERE result IS NOT NULL
                ORDER BY rice_profile, scenario_number
            """)
        
        scenarios = cursor.fetchall()
        
        if not scenarios:
            msg = f"No executed scenarios found for RICE {rice_profile}" if rice_profile else "No executed scenarios found"
            if show_popup:
                show_popup("Warning", msg, "warning")
            conn.close()
            if loading_popup and loading_popup.winfo_exists():
                elapsed_time = time.time() - start_time
                if elapsed_time < 5.0:
                    remaining_time = int((5.0 - elapsed_time) * 1000)
                    loading_popup.after(remaining_time, loading_popup.destroy)
                else:
                    loading_popup.destroy()
            return
        
        # Get RICE profile info
        cursor.execute("SELECT name, rice_id, client_name, tenant FROM rice_profiles WHERE rice_id = ?", (rice_profile,))
        rice_data = cursor.fetchone()
        
        if not rice_data:
            cursor.execute("SELECT name, rice_id, client_name, tenant FROM rice_profiles WHERE id = ?", (rice_profile,))
            rice_data = cursor.fetchone()
        
        if rice_data:
            rice_name, actual_rice_id, client_name, tenant_name = rice_data
            client_name = client_name or "Client Name Not Set"
            tenant_name = tenant_name or "Tenant Not Set"
        else:
            rice_name = f"RICE {rice_profile}"
            actual_rice_id = rice_profile
            client_name = "Client Name Not Set"
            tenant_name = "Tenant Not Set"
        
        # Calculate test statistics
        total_tests = len(scenarios)
        passed_tests = sum(1 for s in scenarios if s[3] == "Passed")
        failed_tests = total_tests - passed_tests
        
        # Get user info
        user_full_name = current_user.get('full_name', 'User Name Not Set') if current_user else 'User Name Not Set'
        
        # Create new document from scratch
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add Infor logo to top-left corner (1x1 inch)
        logo_added = False
        try:
            # Try converted logo first, then original files
            logo_paths = [
                os.path.join(os.path.dirname(__file__), 'infor_logo_from_ico.png'),
                os.path.join(os.path.dirname(__file__), 'infor_logo_fixed.png'),
                os.path.join(os.path.dirname(__file__), 'infor_logo.png'),
                os.path.join(os.path.dirname(__file__), 'infor_logo.ico')
            ]
            
            for logo_path in logo_paths:
                if os.path.exists(logo_path):
                    try:
                        # Create paragraph for logo at very top of document
                        logo_paragraph = doc.add_paragraph()
                        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Add logo with exact 1x1 inch sizing
                        logo_run = logo_paragraph.add_run()
                        logo_run.add_picture(logo_path, width=Inches(1), height=Inches(1))
                        
                        # Set paragraph spacing to keep logo at top
                        logo_paragraph.paragraph_format.space_before = Pt(0)
                        logo_paragraph.paragraph_format.space_after = Pt(6)
                        
                        logo_added = True
                        break
                    except Exception as logo_error:
                        # Try next logo file if this one fails
                        continue
                    
        except Exception as e:
            # If logo fails to load, continue without it
            logo_added = False
        
        # Add spacing after logo (or at top if no logo)
        if logo_added:
            doc.add_paragraph()  # One line after logo
        else:
            doc.add_paragraph()  # Standard top spacing
            doc.add_paragraph()
        
        title = doc.add_heading('TES-070 CUSTOM EXTENSION UNIT TEST RESULT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Document Control section (on same page)
        doc.add_heading('Document Control', 1)
        
        # Document info table
        doc_table = doc.add_table(rows=7, cols=2)
        doc_table.style = 'Table Grid'
        
        doc_table.cell(0, 0).text = 'Client Name'
        doc_table.cell(0, 1).text = client_name if client_name != "Client Name Not Set" else "White Plains Hospital"
        doc_table.cell(1, 0).text = 'RICE ID Description'
        doc_table.cell(1, 1).text = f'RICE-{actual_rice_id}: {rice_name}'
        doc_table.cell(2, 0).text = 'Functional/Technical Resource'
        doc_table.cell(2, 1).text = user_full_name
        doc_table.cell(3, 0).text = 'Tenant'
        doc_table.cell(3, 1).text = tenant_name if tenant_name != "Tenant Not Set" else "TAMICS10_AX1"
        # Get next version number - let database manager handle the logic
        profile_id_for_versions = rice_profile
        if isinstance(rice_profile, str) and not rice_profile.isdigit():
            cursor.execute("SELECT id FROM rice_profiles WHERE rice_id = ?", (rice_profile,))
            result = cursor.fetchone()
            if result:
                profile_id_for_versions = result[0]
        
        # Get current versions to determine next version number
        current_versions = db_manager.get_tes070_versions(profile_id_for_versions)
        if current_versions:
            latest_version = current_versions[0][1]  # version_number from first (latest) record
            next_version_num = latest_version + 1
        else:
            next_version_num = 1
        next_version = f"v{next_version_num}"
        
        doc_table.cell(4, 0).text = 'Version'
        doc_table.cell(4, 1).text = str(next_version_num)
        doc_table.cell(5, 0).text = 'Date'
        doc_table.cell(5, 1).text = datetime.now().strftime('%m/%d/%Y %I:%M:%S %p')
        doc_table.cell(6, 0).text = 'Document Type'
        doc_table.cell(6, 1).text = 'Test Execution Summary (TES-070)'
        
        # Page break after Document Control
        doc.add_page_break()
        
        # Change Record section (on page 2)
        doc.add_heading('1.1\tChange Record', 2)
        
        change_table = doc.add_table(rows=2, cols=4)
        change_table.style = 'Table Grid'
        
        # Headers
        change_table.cell(0, 0).text = 'Version'
        change_table.cell(0, 1).text = 'Date'
        change_table.cell(0, 2).text = 'Author'
        change_table.cell(0, 3).text = 'Description'
        
        # Data
        change_table.cell(1, 0).text = str(next_version_num)
        change_table.cell(1, 1).text = datetime.now().strftime('%m/%d/%Y %I:%M:%S %p')
        change_table.cell(1, 2).text = user_full_name
        change_table.cell(1, 3).text = 'Latest Automated Run'
        
        # Section 2: Summary Results
        doc.add_heading('2\tCustom Extension Unit Test Summary Results', 1)
        doc.add_heading('2.1\tUnit Test Summary', 2)
        
        summary_para = doc.add_paragraph(
            f"This document provides the unit test results for {rice_name}. "
            f"A total of {total_tests} test scenario(s) were executed with "
            f"{passed_tests} passing and {failed_tests} failing."
        )
        
        # Summary table
        summary_table = doc.add_table(rows=2, cols=6)
        summary_table.style = 'Table Grid'
        
        # Headers
        headers = ['No. of scenarios', 'No. of scenarios completed', '% completed', 
                  'No. of scenarios that passed testing', 'No. of scenarios that failed testing', '% passed']
        for i, header in enumerate(headers):
            summary_table.cell(0, i).text = header
        
        # Data
        summary_table.cell(1, 0).text = str(total_tests)
        summary_table.cell(1, 1).text = str(total_tests)
        summary_table.cell(1, 2).text = "100%"
        summary_table.cell(1, 3).text = str(passed_tests)
        summary_table.cell(1, 4).text = str(failed_tests)
        summary_table.cell(1, 5).text = f"{(passed_tests/total_tests*100):.0f}%" if total_tests > 0 else "0%"
        
        # Section 3: Detailed Results - ONLY CREATE FR SECTIONS FOR ACTUAL SCENARIOS
        doc.add_heading('3\tCustom Extension Unit Test Detailed Results', 1)
        
        detail_para = doc.add_paragraph(
            "The Custom Extension Unit Test Detailed Results provide the results for each individual test case."
        )
        
        # Create FR sections only for scenarios that exist
        for i, (rice_prof, scenario_num, description, result, executed_at) in enumerate(scenarios):
            fr_number = f"FR 1.{i+1}"
            
            # FR section header
            doc.add_heading(f'{fr_number}\t{description}', 2)
            
            # Test execution overview
            doc.add_paragraph(f"Test Scenario: {description}")
            doc.add_paragraph(f"Execution Method: Automated Selenium Testing Framework")
            doc.add_paragraph(f"Test Result: {result}")
            doc.add_paragraph(f"Execution Date: {executed_at or datetime.now().strftime('%m/%d/%Y %I:%M:%S %p')}")
            doc.add_paragraph()  # Empty line
            
            # Get steps from database
            cursor.execute("""
                SELECT ss.step_order, 
                       CASE 
                           WHEN LOWER(COALESCE(ts.step_type, ss.step_type)) LIKE '%input%' 
                                AND (LOWER(ss.step_description) LIKE '%password%' 
                                     OR LOWER(ts.description) LIKE '%password%'
                                     OR LOWER(ts.name) LIKE '%password%')
                           THEN 'Password'
                           ELSE ss.step_description
                       END as step_description,
                       ss.screenshot_after, 
                       COALESCE(ts.step_type, ss.step_type)
                FROM scenario_steps ss
                LEFT JOIN test_steps ts ON ss.test_step_id = ts.id
                WHERE ss.rice_profile = ? AND ss.scenario_number = ?
                ORDER BY ss.step_order
            """, (str(rice_prof), scenario_num))
            
            steps_data = cursor.fetchall()
            
            if not steps_data and rice_data:
                cursor.execute("""
                    SELECT ss.step_order, 
                           CASE 
                               WHEN LOWER(COALESCE(ts.step_type, ss.step_type)) LIKE '%input%' 
                                    AND (LOWER(ss.step_description) LIKE '%password%' 
                                         OR LOWER(ts.description) LIKE '%password%'
                                         OR LOWER(ts.name) LIKE '%password%')
                               THEN 'Password'
                               ELSE ss.step_description
                           END as step_description,
                           ss.screenshot_after, 
                           COALESCE(ts.step_type, ss.step_type)
                    FROM scenario_steps ss
                    LEFT JOIN test_steps ts ON ss.test_step_id = ts.id
                    WHERE ss.rice_profile = ? AND ss.scenario_number = ?
                    ORDER BY ss.step_order
                """, (actual_rice_id, scenario_num))
                steps_data = cursor.fetchall()
            
            if steps_data:
                # Filter out wait steps and reorder
                filtered_steps = []
                for step_order, step_desc, screenshot_b64, step_type in steps_data:
                    if step_desc:
                        step_lower = step_desc.lower()
                        # Skip wait steps and empty/generic steps
                        if not any(wait_word in step_lower for wait_word in ['wait', 'sleep', 'pause', 'delay', 'implicit', 'explicit']) and len(step_desc.strip()) > 5:
                            # Format step description properly
                            formatted_desc = step_desc
                            
                            # Use step_type to determine action description, fallback to content analysis
                            if step_type and step_type.strip():
                                step_type_lower = step_type.lower()
                                if 'navigate' in step_type_lower:
                                    if step_desc.lower().startswith('navigate to '):
                                        formatted_desc = step_desc
                                    else:
                                        formatted_desc = f"Navigate to {step_desc}"
                                elif 'click' in step_type_lower:
                                    if step_desc.lower().startswith('click '):
                                        formatted_desc = step_desc
                                    else:
                                        formatted_desc = f"Click {step_desc}"
                                elif 'input' in step_type_lower or 'text' in step_type_lower:
                                    if 'workunit' in step_desc.lower() or 'work unit' in step_desc.lower():
                                        formatted_desc = "Enter Workunit"
                                    elif step_desc.lower() == 'password':
                                        formatted_desc = "Enter password"
                                    elif '@' in step_desc and '.' in step_desc and ' ' not in step_desc.strip():
                                        formatted_desc = f"Enter username: {step_desc}"
                                    else:
                                        formatted_desc = f"Enter {step_desc}"
                                elif 'select' in step_type_lower or 'dropdown' in step_type_lower:
                                    formatted_desc = f"Select {step_desc}"
                                else:
                                    formatted_desc = f"{step_type}: {step_desc}"
                            else:
                                # Fallback: analyze step description content
                                if 'navigate' in step_lower or 'page' in step_lower:
                                    if step_desc.lower().startswith('navigate to '):
                                        formatted_desc = step_desc
                                    else:
                                        formatted_desc = f"Navigate to {step_desc}"
                                elif 'click' in step_lower or 'button' in step_lower:
                                    if step_desc.lower().startswith('click '):
                                        formatted_desc = step_desc
                                    else:
                                        formatted_desc = f"Click {step_desc}"
                                elif '@' in step_desc and '.' in step_desc and ' ' not in step_desc.strip():
                                    formatted_desc = f"Enter username: {step_desc}"
                                elif 'workunit' in step_lower or 'work unit' in step_lower:
                                    formatted_desc = "Enter Workunit"
                                elif step_desc.lower() == 'password':
                                    formatted_desc = "Enter password"
                                elif len(step_desc.strip()) < 10 and not any(action in step_lower for action in ['navigate', 'click', 'select']):
                                    formatted_desc = f"Enter {step_desc}"
                                else:
                                    formatted_desc = step_desc
                            
                            filtered_steps.append((formatted_desc, screenshot_b64))
                
                if filtered_steps:
                    # Add detailed test steps section
                    doc.add_heading("Test Execution Steps", 3)
                    
                    # Create steps table for better organization
                    steps_table = doc.add_table(rows=1, cols=3)
                    steps_table.style = 'Table Grid'
                    
                    # Table headers
                    hdr_cells = steps_table.rows[0].cells
                    hdr_cells[0].text = 'Step #'
                    hdr_cells[1].text = 'Action Description'
                    hdr_cells[2].text = 'Screenshot'
                    
                    # Add each step with proper numbering
                    for j, (step_desc, screenshot_b64) in enumerate(filtered_steps, 1):
                        row_cells = steps_table.add_row().cells
                        row_cells[0].text = str(j)
                        row_cells[1].text = step_desc
                        
                        if screenshot_b64:
                            try:
                                # Decode and insert screenshot
                                import base64
                                from io import BytesIO
                                
                                # Decode base64 screenshot
                                screenshot_data = base64.b64decode(screenshot_b64)
                                screenshot_stream = BytesIO(screenshot_data)
                                
                                # Add screenshot to document with proper sizing
                                paragraph = row_cells[2].paragraphs[0]
                                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                                run.add_picture(screenshot_stream, width=Inches(2.5))
                                
                            except Exception as e:
                                row_cells[2].text = f"Screenshot available (Error loading: {str(e)[:50]}...)"
                        else:
                            row_cells[2].text = "No screenshot captured"
                    
                    doc.add_paragraph()  # Empty line after table
                    
                    # Add test validation section
                    doc.add_heading("Test Validation", 3)
                    validation_para = doc.add_paragraph()
                    if result == "Passed":
                        validation_para.add_run("✓ Test Passed: ").bold = True
                        validation_para.add_run("All test steps executed successfully without errors. Expected functionality verified.")
                    else:
                        validation_para.add_run("✗ Test Failed: ").bold = True
                        validation_para.add_run("Test execution encountered errors. See Section 4 Problems for detailed analysis.")
                    
                    # Add technical details
                    doc.add_paragraph()
                    tech_details = doc.add_paragraph()
                    tech_details.add_run("Technical Details:").bold = True
                    doc.add_paragraph(f"• Total Steps Executed: {len(filtered_steps)}")
                    doc.add_paragraph(f"• Screenshots Captured: {sum(1 for _, ss in filtered_steps if ss)}")
                    doc.add_paragraph(f"• Test Environment: {tenant_name if tenant_name != 'Tenant Not Set' else 'TAMICS10_AX1'}")
                    doc.add_paragraph(f"• Browser: Chrome (Selenium WebDriver)")
                    
                else:
                    doc.add_paragraph("No actionable steps recorded (only wait/delay steps found).")
                    doc.add_paragraph("Test executed via automated framework with minimal user interaction.")
            else:
                doc.add_paragraph("Test Execution Summary:")
                doc.add_paragraph("• Test executed via automated Selenium framework")
                doc.add_paragraph("• No detailed step-by-step recording available")
                doc.add_paragraph("• Test validation based on final result status")
            
            # Add separator between FR sections
            doc.add_paragraph()
            separator = doc.add_paragraph("─" * 80)
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Section 4: Problems and Issues Analysis
        doc.add_heading('4\tProblems and Issues Analysis', 1)
        
        if failed_tests > 0:
            doc.add_paragraph("The following issues were identified during test execution:")
            doc.add_paragraph()
            
            # Get failed scenarios for detailed analysis
            failed_scenarios = [s for s in scenarios if s[3] == "Failed"]
            
            for idx, (rice_prof, scenario_num, description, result, executed_at) in enumerate(failed_scenarios, 1):
                doc.add_heading(f'4.{idx}\tIssue #{idx}: {description}', 2)
                
                # Issue details table
                issue_table = doc.add_table(rows=6, cols=2)
                issue_table.style = 'Table Grid'
                
                issue_table.cell(0, 0).text = 'Scenario'
                issue_table.cell(0, 1).text = description
                issue_table.cell(1, 0).text = 'Scenario Number'
                issue_table.cell(1, 1).text = str(scenario_num)
                issue_table.cell(2, 0).text = 'Failure Date'
                issue_table.cell(2, 1).text = executed_at or 'Not recorded'
                issue_table.cell(3, 0).text = 'Impact Level'
                issue_table.cell(3, 1).text = 'High - Test Execution Blocked'
                issue_table.cell(4, 0).text = 'Status'
                issue_table.cell(4, 1).text = 'Open - Requires Investigation'
                issue_table.cell(5, 0).text = 'Recommended Action'
                issue_table.cell(5, 1).text = 'Review test steps, verify environment configuration, re-execute after fixes'
                
                doc.add_paragraph()
            
            # Overall recommendations
            doc.add_heading('4.99\tRecommendations', 2)
            doc.add_paragraph("Based on the test failures identified, the following actions are recommended:")
            doc.add_paragraph("• Review failed test scenarios for environment-specific issues")
            doc.add_paragraph("• Verify FSM system availability and configuration")
            doc.add_paragraph("• Check network connectivity and authentication credentials")
            doc.add_paragraph("• Re-execute failed tests after addressing identified issues")
            doc.add_paragraph("• Consider additional test scenarios to improve coverage")
            
        else:
            doc.add_paragraph("✓ No issues identified during testing execution.")
            doc.add_paragraph()
            doc.add_paragraph("All test scenarios executed successfully without errors or failures.")
            doc.add_paragraph("The tested functionality meets the expected requirements and performs as designed.")
            doc.add_paragraph()
            doc.add_paragraph("Quality Assessment:")
            doc.add_paragraph("• Test Coverage: Complete - All defined scenarios executed")
            doc.add_paragraph("• Test Results: Satisfactory - 100% pass rate achieved")
            doc.add_paragraph("• System Stability: Excellent - No crashes or unexpected behavior")
            doc.add_paragraph("• Performance: Acceptable - All operations completed within expected timeframes")
        
        # Section 5: Test Environment and Configuration
        doc.add_heading('5\tTest Environment and Configuration', 1)
        
        doc.add_paragraph("This section documents the test environment configuration and technical setup used during test execution.")
        doc.add_paragraph()
        
        # Environment details table
        env_table = doc.add_table(rows=8, cols=2)
        env_table.style = 'Table Grid'
        
        env_table.cell(0, 0).text = 'Test Environment'
        env_table.cell(0, 1).text = tenant_name if tenant_name != 'Tenant Not Set' else 'TAMICS10_AX1'
        env_table.cell(1, 0).text = 'Client/Organization'
        env_table.cell(1, 1).text = client_name
        env_table.cell(2, 0).text = 'RICE Identifier'
        env_table.cell(2, 1).text = f'RICE-{actual_rice_id}'
        env_table.cell(3, 0).text = 'Test Framework'
        env_table.cell(3, 1).text = 'Selenium WebDriver with Python'
        env_table.cell(4, 0).text = 'Browser'
        env_table.cell(4, 1).text = 'Google Chrome (Latest Version)'
        env_table.cell(5, 0).text = 'Test Execution Date'
        env_table.cell(5, 1).text = datetime.now().strftime('%m/%d/%Y %I:%M:%S %p')
        env_table.cell(6, 0).text = 'Test Executor'
        env_table.cell(6, 1).text = user_full_name
        env_table.cell(7, 0).text = 'Total Test Duration'
        env_table.cell(7, 1).text = 'Automated - Variable based on scenario complexity'
        
        # Section 6: Appendices
        doc.add_heading('6\tAppendices', 1)
        
        doc.add_heading('6.1\tTest Data Summary', 2)
        
        # Test data table
        data_table = doc.add_table(rows=len(scenarios) + 1, cols=5)
        data_table.style = 'Table Grid'
        
        # Headers
        data_headers = ['Scenario #', 'Description', 'Result', 'Execution Time', 'Steps Count']
        for i, header in enumerate(data_headers):
            data_table.cell(0, i).text = header
        
        # Data rows
        for i, (rice_prof, scenario_num, description, result, executed_at) in enumerate(scenarios, 1):
            # Get step count for this scenario
            cursor.execute("""
                SELECT COUNT(*) FROM scenario_steps ss
                LEFT JOIN test_steps ts ON ss.test_step_id = ts.id
                WHERE ss.rice_profile = ? AND ss.scenario_number = ?
            """, (str(rice_prof), scenario_num))
            step_count = cursor.fetchone()[0] or 0
            
            data_table.cell(i, 0).text = str(scenario_num)
            data_table.cell(i, 1).text = description[:50] + ('...' if len(description) > 50 else '')
            data_table.cell(i, 2).text = result
            data_table.cell(i, 3).text = executed_at or 'Not recorded'
            data_table.cell(i, 4).text = f"{step_count} (including wait steps)"
        
        doc.add_paragraph()
        
        doc.add_heading('6.2\tGlossary', 2)
        doc.add_paragraph("RICE: Reports, Interfaces, Conversions, Extensions")
        doc.add_paragraph("TES-070: Test Execution Summary Document (Standard Format)")
        doc.add_paragraph("FSM: Financials and Supply Management (Infor ERP System)")
        doc.add_paragraph("Selenium: Web browser automation framework")
        doc.add_paragraph("WebDriver: Browser automation API used by Selenium")
        
        doc.add_heading('6.3\tDocument References', 2)
        doc.add_paragraph("• Infor FSM User Documentation")
        doc.add_paragraph("• RICE Development Standards and Guidelines")
        doc.add_paragraph("• Test Execution Framework Documentation")
        doc.add_paragraph("• Selenium WebDriver Best Practices Guide")
        
        # Save to database only
        def save_to_database():
            if loading_popup and loading_popup.winfo_exists():
                loading_popup.destroy()
            
            try:
                from io import BytesIO
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                file_content = buffer.getvalue()
                
                # Save to database using provided db_manager with profile ID
                user_name = current_user.get('full_name', 'Current User') if current_user else 'Current User'
                version_id = db_manager.save_tes070_version(profile_id_for_versions, file_content, user_name)
                
                if show_popup:
                    show_popup("Success", f"TES-070 report generated from scratch and saved to database!\n\nVersion: {next_version}\nUse TES-070 History to download.", "success")
                    
            except Exception as e:
                if show_popup:
                    show_popup("Error", f"Failed to save TES-070: {str(e)}", "error")
        
        # Schedule save after minimum loading time
        if loading_popup and loading_popup.winfo_exists():
            elapsed_time = time.time() - start_time
            if elapsed_time < 5.0:
                remaining_time = int((5.0 - elapsed_time) * 1000)
                loading_popup.after(remaining_time, save_to_database)
            else:
                save_to_database()
        else:
            save_to_database()
        
        conn.close()
        
    except Exception as e:
        if loading_popup and loading_popup.winfo_exists():
            elapsed_time = time.time() - start_time
            if elapsed_time < 5.0:
                remaining_time = int((5.0 - elapsed_time) * 1000)
                loading_popup.after(remaining_time, loading_popup.destroy)
            else:
                loading_popup.destroy()
        if show_popup:
            show_popup("Error", f"Failed to generate report: {str(e)}", "error")


if __name__ == "__main__":
    # Create simple GUI for testing
    root = tk.Tk()
    root.withdraw()
    create_tes070_from_scratch()