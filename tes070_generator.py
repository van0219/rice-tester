#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sqlite3
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog
from database_manager import DatabaseManager
from rice_dialogs import center_dialog
from io import BytesIO

def generate_tes070_report(rice_profile, show_popup=None, current_user=None, db_manager=None):
    """Generate TES-070 and save to database only (no file download)"""
    
    # Always define start_time at function start
    start_time = time.time()
    
    # Show loading dialog
    loading_popup = None
    if show_popup:
        
        # Create loading dialog
        loading_popup = tk.Toplevel()
        loading_popup.title("Generating TES-070")
        center_dialog(loading_popup, 400, 236)
        loading_popup.configure(bg='#ffffff')
        loading_popup.resizable(False, False)
        loading_popup.grab_set()
        
        try:
            loading_popup.iconbitmap("infor_logo.ico")
        except:
            pass
        
        # Header
        header_frame = tk.Frame(loading_popup, bg='#3b82f6', height=60)
        header_frame.pack(fill="x")
        header_frame.pack_propagate(False)
        
        header_label = tk.Label(header_frame, text="📄 Generating TES-070 Report", 
                               font=('Segoe UI', 14, 'bold'), bg='#3b82f6', fg='#ffffff')
        header_label.pack(expand=True)
        
        # Content
        content_frame = tk.Frame(loading_popup, bg='#ffffff', padx=20, pady=20)
        content_frame.pack(fill="both", expand=True)
        
        # Loading message
        loading_label = tk.Label(content_frame, text="Please wait while generating your TES-070 report...\n\nSaving to database for version control.", 
                                font=('Segoe UI', 10), bg='#ffffff', justify="center")
        loading_label.pack(pady=(10, 20))
        
        # Progress indicator
        progress_label = tk.Label(content_frame, text="●●●", 
                                 font=('Segoe UI', 16), bg='#ffffff', fg='#3b82f6')
        progress_label.pack()
        
        # Animate the progress dots
        def animate_dots():
            current = progress_label.cget('text')
            if current == "●●●":
                progress_label.config(text="○●●")
            elif current == "○●●":
                progress_label.config(text="○○●")
            elif current == "○○●":
                progress_label.config(text="○○○")
            else:
                progress_label.config(text="●●●")
            
            if loading_popup.winfo_exists():
                loading_popup.after(300, animate_dots)
        
        # Start animation
        animate_dots()
        loading_popup.update()
    
    try:
        # Database path
        db_path = r"d:\AmazonQ\InforQ\RICE_Tester\fsm_tester.db"
        template_path = r"d:\AmazonQ\InforQ\RICE_Tester\TES-070-Template\TES-070_Custom_Extension_Unit_Test_Results_v2.0.docx"
    
        if not os.path.exists(db_path):
            if loading_popup and loading_popup.winfo_exists():
                elapsed_time = time.time() - start_time
                if elapsed_time < 5.0:
                    remaining_time = int((5.0 - elapsed_time) * 1000)
                    loading_popup.after(remaining_time, loading_popup.destroy)
                else:
                    loading_popup.destroy()
            if show_popup:
                show_popup("Error", "Database not found", "error")
            return
            
        if not os.path.exists(template_path):
            if loading_popup and loading_popup.winfo_exists():
                elapsed_time = time.time() - start_time
                if elapsed_time < 5.0:
                    remaining_time = int((5.0 - elapsed_time) * 1000)
                    loading_popup.after(remaining_time, loading_popup.destroy)
                else:
                    loading_popup.destroy()
            if show_popup:
                show_popup("Error", "TES-070 template not found", "error")
            return
        
        # Connect to database
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Check for "not run" scenarios
        if rice_profile:
            cursor.execute("""
                SELECT COUNT(*) FROM scenarios s
                JOIN rice_profiles rp ON s.rice_profile = rp.id
                WHERE (rp.id = ? OR rp.rice_id = ?) AND (s.result IS NULL OR s.result = 'Not run')
            """, (rice_profile, rice_profile))
            not_run_count = cursor.fetchone()[0]
            
            if not_run_count > 0:
                cursor.execute("SELECT rice_id FROM rice_profiles WHERE id = ? OR rice_id = ?", (rice_profile, rice_profile))
                rice_id_result = cursor.fetchone()
                display_rice_id = rice_id_result[0] if rice_id_result else rice_profile
                
                if show_popup:
                    # Create custom dialog with increased height (0.5 inch = 36px)
                    error_popup = tk.Toplevel()
                    error_popup.title("Cannot Generate TES-070")
                    center_dialog(error_popup, 450, 272)  # Increased from 236 to 272 (36px more)
                    error_popup.configure(bg='#ffffff')
                    error_popup.resizable(False, False)
                    error_popup.grab_set()
                    
                    try:
                        error_popup.iconbitmap("infor_logo.ico")
                    except:
                        pass
                    
                    # Header
                    header_frame = tk.Frame(error_popup, bg='#ef4444', height=60)
                    header_frame.pack(fill="x")
                    header_frame.pack_propagate(False)
                    
                    tk.Label(header_frame, text="⚠️ Cannot Generate TES-070", 
                            font=('Segoe UI', 14, 'bold'), bg='#ef4444', fg='#ffffff').pack(expand=True)
                    
                    # Content
                    content_frame = tk.Frame(error_popup, bg='#ffffff', padx=20, pady=20)
                    content_frame.pack(fill="both", expand=True)
                    
                    message_text = (f"Cannot generate TES-070 report for RICE {display_rice_id}.\n\n"
                                  f"Found {not_run_count} scenario(s) with 'Not run' status.\n\n"
                                  "Please execute all scenarios before generating the report.")
                    
                    tk.Label(content_frame, text=message_text, font=('Segoe UI', 10), 
                            bg='#ffffff', fg='#374151', justify="center").pack(pady=(0, 20))
                    
                    # OK button
                    tk.Button(content_frame, text="OK", font=('Segoe UI', 10, 'bold'), 
                             bg='#ef4444', fg='#ffffff', relief='flat', padx=20, pady=8, 
                             cursor='hand2', bd=0, command=error_popup.destroy).pack()
                    
                    error_popup.focus_set()
                conn.close()
                if loading_popup and loading_popup.winfo_exists():
                    elapsed_time = time.time() - start_time
                    if elapsed_time < 5.0:
                        remaining_time = int((5.0 - elapsed_time) * 1000)
                        loading_popup.after(remaining_time, loading_popup.destroy)
                    else:
                        loading_popup.destroy()
                return
        
        # Get scenario data
        if rice_profile:
            cursor.execute("""
                SELECT s.rice_profile, s.scenario_number, s.description, s.result, s.executed_at
                FROM scenarios s
                JOIN rice_profiles rp ON s.rice_profile = rp.id
                WHERE s.result IS NOT NULL AND (rp.id = ? OR rp.rice_id = ?)
                ORDER BY s.scenario_number
            """, (rice_profile, rice_profile))
        else:
            cursor.execute("""
                SELECT rice_profile, scenario_number, description, result, executed_at
                FROM scenarios 
                WHERE result IS NOT NULL
                ORDER BY rice_profile, scenario_number
            """)
        
        scenarios = cursor.fetchall()
        
        if not scenarios:
            msg = f"No executed scenarios found for RICE {rice_profile}" if rice_profile else "No executed scenarios found"
            if show_popup:
                show_popup("Warning", msg, "warning")
            conn.close()
            if loading_popup and loading_popup.winfo_exists():
                elapsed_time = time.time() - start_time
                if elapsed_time < 5.0:
                    remaining_time = int((5.0 - elapsed_time) * 1000)
                    loading_popup.after(remaining_time, loading_popup.destroy)
                else:
                    loading_popup.destroy()
            return
        
        # Calculate test statistics
        total_tests = len(scenarios)
        passed_tests = sum(1 for s in scenarios if s[3] == "Passed")
        failed_tests = total_tests - passed_tests
        
        # Load template
        doc = Document(template_path)
        
        # Get RICE profile info
        cursor.execute("SELECT name, rice_id, client_name, tenant FROM rice_profiles WHERE rice_id = ?", (rice_profile,))
        rice_data = cursor.fetchone()
        
        if not rice_data:
            cursor.execute("SELECT name, rice_id, client_name, tenant FROM rice_profiles WHERE id = ?", (rice_profile,))
            rice_data = cursor.fetchone()
        
        if rice_data:
            rice_name, actual_rice_id, client_name, tenant_name = rice_data
            client_name = client_name or "Client Name Not Set"
            tenant_name = tenant_name or "Tenant Not Set"
        else:
            rice_name = f"RICE {rice_profile}"
            actual_rice_id = rice_profile
            client_name = "Client Name Not Set"
            tenant_name = "Tenant Not Set"
        
        # Replace placeholders in document (same logic as original function)
        user_full_name = current_user.get('full_name', 'User Name Not Set') if current_user else 'User Name Not Set'
        
        # First pass: Update all FR titles with actual scenario descriptions (dynamic numbering)
        for i, para in enumerate(doc.paragraphs):
            if para.text.startswith('FR ') and any(x in para.text for x in ['1.1', '1.2', '1.3']):
                # Extract FR number and convert to scenario index
                scenario_index = None
                if '1.1' in para.text:
                    scenario_index = 0
                elif '1.2' in para.text:
                    scenario_index = 1
                elif '1.3' in para.text:
                    scenario_index = 2
                
                if scenario_index is not None and scenario_index < len(scenarios):
                    # Dynamic FR numbering: FR 1.1, FR 1.2, FR 1.3 based on scenario count
                    fr_number = f'FR 1.{scenario_index + 1}'
                    scenario_desc = scenarios[scenario_index][2]  # Get scenario description
                    para.text = f'{fr_number}\t{scenario_desc}'
        
        for para in doc.paragraphs:
            if '<Client Name>' in para.text:
                para.text = para.text.replace('<Client Name>', client_name)
            if '<RICE ID Description>' in para.text:
                para.text = para.text.replace('<RICE ID Description>', f'RICE-{actual_rice_id}: {rice_name}')
            if '<Requirement Definition Narrative>' in para.text:
                # Check if this is an FR section and use scenario description
                if para.text.startswith('FR '):
                    for num in ['1.1', '1.2', '1.3']:
                        if num in para.text:
                            scenario_index = int(num.split('.')[1]) - 1
                            if scenario_index < len(scenarios):
                                scenario_desc = scenarios[scenario_index][2]
                                para.text = para.text.replace('<Requirement Definition Narrative>', scenario_desc)
                            break
                else:
                    para.text = para.text.replace('<Requirement Definition Narrative>', f'Automated testing for {rice_name}')

            if '<Scenario Description>' in para.text:
                para.text = para.text.replace('<Scenario Description>', 'Automated test scenario execution')
            if '<Passed/Failed>' in para.text:
                para.text = para.text.replace('<Passed/Failed>', 'See test results table below')
            if '<Screenshots and navigation>' in para.text:
                para.text = para.text.replace('<Screenshots and navigation>', 'Screenshots captured during automated execution')
            if '<Issue Description>' in para.text:
                para.text = para.text.replace('<Issue Description>', 'No issues identified during testing')
            if '<Description of proposed resolution>' in para.text:
                para.text = para.text.replace('<Description of proposed resolution>', 'N/A - All tests passed')
            if '<Author>' in para.text:
                para.text = para.text.replace('<Author>', user_full_name)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Functional/Technical Resource>' in cell.text:
                        user_full_name = current_user.get('full_name', 'User Name Not Set') if current_user else 'User Name Not Set'
                        cell.text = cell.text.replace('<Functional/Technical Resource>', user_full_name)
                    if '<Tenant>' in cell.text:
                        cell.text = cell.text.replace('<Tenant>', tenant_name)
                    if '<Version>' in cell.text:
                        cell.text = cell.text.replace('<Version>', '1.0')
                    if 'MM/dd/yyyy' in cell.text:
                        current_date = datetime.now().strftime('%m/%d/%Y')
                        cell.text = cell.text.replace('MM/dd/yyyy', current_date)
                    if '<No. of scenarios>' in cell.text:
                        cell.text = cell.text.replace('<No. of scenarios>', str(total_tests))
                    if '<No. of scenarios completed>' in cell.text:
                        cell.text = cell.text.replace('<No. of scenarios completed>', str(total_tests))
                    if '<No. of scenarios that passed testing>' in cell.text:
                        cell.text = cell.text.replace('<No. of scenarios that passed testing>', str(passed_tests))
                    if '<No. of scenarios that failed testing>' in cell.text:
                        cell.text = cell.text.replace('<No. of scenarios that failed testing>', str(failed_tests))
        
        # Update summary table
        summary_table = doc.tables[4]
        summary_row = summary_table.rows[1]
        summary_row.cells[0].text = str(total_tests)
        summary_row.cells[1].text = str(total_tests)
        summary_row.cells[2].text = "100%"
        summary_row.cells[3].text = str(passed_tests)
        summary_row.cells[4].text = str(failed_tests)
        summary_row.cells[5].text = f"{(passed_tests/total_tests*100):.0f}%" if total_tests > 0 else "0%"
        
        # FINAL FIX: Process document systematically - populate ALL FR sections
        
        # Step 1: Remove unused FR sections FIRST (before any content population)
        if len(scenarios) < 3:
            # Simple approach: Remove <Steps and Screenshots> placeholders for unused FR sections
            steps_placeholders_to_remove = []
            if len(scenarios) == 1:
                # Keep only first <Steps and Screenshots>, remove 2nd and 3rd
                steps_placeholders_to_remove = [1, 2]  # Remove 2nd and 3rd occurrences
            elif len(scenarios) == 2:
                # Keep first two <Steps and Screenshots>, remove 3rd
                steps_placeholders_to_remove = [2]  # Remove 3rd occurrence
            
            # Find and remove unused <Steps and Screenshots> sections
            steps_count = 0
            paragraphs_to_remove = []
            
            for i, para in enumerate(doc.paragraphs):
                if '<Steps and Screenshots>' in para.text:
                    if steps_count in steps_placeholders_to_remove:
                        # Mark this paragraph and surrounding content for removal
                        # Look backwards to find the FR header
                        start_idx = i
                        for j in range(i-1, -1, -1):
                            if doc.paragraphs[j].text.startswith('FR 1.'):
                                start_idx = j
                                break
                        
                        # Look forwards to find the end (next FR section or section 4)
                        end_idx = len(doc.paragraphs)
                        for j in range(i+1, len(doc.paragraphs)):
                            if (doc.paragraphs[j].text.startswith('FR 1.') or 
                                doc.paragraphs[j].text.startswith('4\t')):
                                end_idx = j
                                break
                        
                        # Mark all paragraphs in this section for removal
                        for k in range(start_idx, end_idx):
                            paragraphs_to_remove.append(k)
                    
                    steps_count += 1
            
            # Remove paragraphs in reverse order to maintain indices
            for para_index in reversed(sorted(set(paragraphs_to_remove))):
                if para_index < len(doc.paragraphs):
                    p = doc.paragraphs[para_index]
                    p._element.getparent().remove(p._element)
        
        # Step 2: Replace all "Automated testing for Asset Accounts Extract" with scenario descriptions
        for para in doc.paragraphs:
            if 'Automated testing for Asset Accounts Extract' in para.text:
                if scenarios:
                    para.text = scenarios[0][2]  # Use first scenario description
        
        # Step 3: Replace <Steps and Screenshots> with actual steps - SIMPLE APPROACH
        steps_replacements = {}
        
        # Pre-calculate steps for each scenario
        for i, (rice_prof, scenario_num, description, result, executed_at) in enumerate(scenarios):
            # Get steps from database
            cursor.execute("""
                SELECT ss.step_order, ss.step_description, ss.screenshot_after
                FROM scenario_steps ss
                WHERE ss.rice_profile = ? AND ss.scenario_number = ?
                ORDER BY ss.step_order
            """, (str(rice_prof), scenario_num))
            
            steps_data = cursor.fetchall()
            
            if not steps_data and rice_data:
                cursor.execute("""
                    SELECT ss.step_order, ss.step_description, ss.screenshot_after
                    FROM scenario_steps ss
                    WHERE ss.rice_profile = ? AND ss.scenario_number = ?
                    ORDER BY ss.step_order
                """, (actual_rice_id, scenario_num))
                steps_data = cursor.fetchall()
            
            if steps_data:
                # Build steps content
                steps_content = "\n\nTEST EXECUTION STEPS:\n\n"
                
                # Filter out wait steps first, then renumber
                filtered_steps = []
                for step_order, step_desc, screenshot_b64 in steps_data:
                    if step_desc and 'wait' not in step_desc.lower():
                        filtered_steps.append((step_desc, screenshot_b64))
                
                # Build text content with proper numbering
                for j, (step_desc, screenshot_b64) in enumerate(filtered_steps, 1):
                    steps_content += f"Step {j}: {step_desc}\n"
                    if screenshot_b64:
                        steps_content += "[Screenshot will be inserted here]\n"
                    steps_content += "\n"
                
                steps_content += f"\nResult: {result}\nExecution Date: {executed_at or 'Not recorded'}"
                steps_replacements[f'FR 1.{i+1}'] = steps_content
            else:
                steps_replacements[f'FR 1.{i+1}'] = '\n\nNo detailed steps recorded for this scenario.\nTest executed via automated framework.'
        
        # Now replace <Steps and Screenshots> based on which FR section we're in
        current_fr = None
        for para in doc.paragraphs:
            # Track current FR section
            if para.text.startswith('FR 1.1'):
                current_fr = 'FR 1.1'
            elif para.text.startswith('FR 1.2'):
                current_fr = 'FR 1.2'
            elif para.text.startswith('FR 1.3'):
                current_fr = 'FR 1.3'
            elif para.text.startswith('4\t'):
                current_fr = None
            
            # Replace placeholder with correct content
            if '<Steps and Screenshots>' in para.text and current_fr:
                if current_fr in steps_replacements:
                    para.text = para.text.replace('<Steps and Screenshots>', steps_replacements[current_fr])
                else:
                    para.text = para.text.replace('<Steps and Screenshots>', '')
        

        
        # Save to database only
        def save_to_database():
            if loading_popup and loading_popup.winfo_exists():
                loading_popup.destroy()
            
            try:
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                file_content = buffer.getvalue()
                
                # Save to database using provided db_manager
                user_name = current_user.get('full_name', 'Current User') if current_user else 'Current User'
                version_id = db_manager.save_tes070_version(rice_profile, file_content, user_name)
                
                if show_popup:
                    show_popup("Success", f"TES-070 report generated and saved to database!\n\nVersion ID: {version_id}\nUse TES-070 History to download.", "success")
                    
            except Exception as e:
                if show_popup:
                    show_popup("Error", f"Failed to save TES-070: {str(e)}", "error")
        
        # Schedule save after minimum loading time
        if loading_popup and loading_popup.winfo_exists():
            elapsed_time = time.time() - start_time
            if elapsed_time < 5.0:
                remaining_time = int((5.0 - elapsed_time) * 1000)
                loading_popup.after(remaining_time, save_to_database)
            else:
                save_to_database()
        else:
            save_to_database()
        
        conn.close()
        
    except Exception as e:
        if loading_popup and loading_popup.winfo_exists():
            elapsed_time = time.time() - start_time
            if elapsed_time < 5.0:
                remaining_time = int((5.0 - elapsed_time) * 1000)
                loading_popup.after(remaining_time, loading_popup.destroy)
            else:
                loading_popup.destroy()
        if show_popup:
            show_popup("Error", f"Failed to generate report: {str(e)}", "error")



if __name__ == "__main__":
    # Create simple GUI for testing
    root = tk.Tk()
    root.withdraw()
    generate_tes070_report()
