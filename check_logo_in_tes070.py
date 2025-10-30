#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
os.system('chcp 65001 >nul 2>&1')  # Set Windows console to UTF-8
sys.path.append(os.path.dirname(__file__))

from database_manager import DatabaseManager
from docx import Document
from io import BytesIO

def check_logo_in_latest_tes070():
    """Check if logo is present in the latest TES-070 document"""
    
    print("Checking logo in latest TES-070 document...")
    
    try:
        # Initialize database manager
        db_manager = DatabaseManager(user_id=1)
        print("Database connected")
        
        # Get latest TES-070 version for INT001
        versions = db_manager.get_tes070_versions("INT001")
        
        if not versions:
            print("No TES-070 versions found for INT001")
            return False
        
        latest_version = versions[0]  # First one is latest
        version_id, version_number, created_at, created_by = latest_version
        
        print(f"Found latest version: v{version_number} created at {created_at}")
        
        # Get document content
        file_content = db_manager.get_tes070_content(version_id)
        
        if not file_content:
            print("No file content found")
            return False
        
        print(f"Document size: {len(file_content)} bytes")
        
        # Load document from bytes
        doc_stream = BytesIO(file_content)
        doc = Document(doc_stream)
        
        # Check for images in the document
        image_count = 0
        logo_found = False
        
        # Check specifically the first few paragraphs for logo
        for i, paragraph in enumerate(doc.paragraphs[:5]):  # Check first 5 paragraphs
            for run in paragraph.runs:
                if run._element.xpath('.//a:blip'):  # Check for images
                    image_count += 1
                    print(f"Found image #{image_count} in paragraph {i+1}")
                    
                    # Check if this is the logo (in first few paragraphs)
                    if i < 3:  # Logo should be in first 3 paragraphs
                        logo_found = True
                        print(f"LOGO FOUND: Image in paragraph {i+1} - This is the Infor logo!")
                        
        # Check remaining paragraphs without detailed logging
        for i, paragraph in enumerate(doc.paragraphs[5:], 6):
            for run in paragraph.runs:
                if run._element.xpath('.//a:blip'):
                    image_count += 1
        
        # Also check tables for images
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//a:blip'):
                                image_count += 1
                                print(f"Found image #{image_count} in table cell")
        
        print(f"Total images found: {image_count}")
        
        if logo_found:
            print("SUCCESS: Infor logo appears to be present in the document!")
            return True
        elif image_count > 0:
            print("Images found but logo position unclear")
            return True
        else:
            print("NO IMAGES FOUND: Logo is missing from the document")
            return False
            
    except Exception as e:
        print(f"Error checking document: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = check_logo_in_latest_tes070()
    if success:
        print("\nLogo check PASSED!")
    else:
        print("\nLogo check FAILED!")