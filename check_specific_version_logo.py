#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
os.system('chcp 65001 >nul 2>&1')  # Set Windows console to UTF-8
sys.path.append(os.path.dirname(__file__))

from database_manager import DatabaseManager
from docx import Document
from io import BytesIO

def check_specific_version_logo(version_id):
    """Check if logo is present in a specific TES-070 version"""
    
    print(f"Checking logo in TES-070 version ID {version_id}...")
    
    try:
        db_manager = DatabaseManager(user_id=1)
        
        # Get document content
        file_content = db_manager.get_tes070_content(version_id)
        
        if not file_content:
            print("No file content found")
            return False
        
        print(f"Document size: {len(file_content)} bytes")
        
        # Load document from bytes
        doc_stream = BytesIO(file_content)
        doc = Document(doc_stream)
        
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        
        # Check first 5 paragraphs specifically for logo
        logo_found = False
        for i, paragraph in enumerate(doc.paragraphs[:5]):
            text = paragraph.text.strip()
            has_images = bool(paragraph._element.xpath('.//a:blip'))
            
            print(f"  Paragraph {i+1}: {'[IMAGE]' if has_images else '[NO IMAGE]'} '{text[:30]}{'...' if len(text) > 30 else ''}'")
            
            if has_images and i < 3:  # Logo should be in first 3 paragraphs
                logo_found = True
                print(f"    *** LOGO FOUND in paragraph {i+1}! ***")
        
        # Count all images
        total_images = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//a:blip'):
                    total_images += 1
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run._element.xpath('.//a:blip'):
                                total_images += 1
        
        print(f"Total images in document: {total_images}")
        
        if logo_found:
            print("SUCCESS: Infor logo found in title area!")
            return True
        else:
            print("ISSUE: No logo found in title area (first 3 paragraphs)")
            return False
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # Check the latest version (ID 33)
    success = check_specific_version_logo(33)
    
    if not success:
        print("\nTrying to check version 29 (Test User version)...")
        success = check_specific_version_logo(29)
    
    if success:
        print("\nLogo check PASSED!")
    else:
        print("\nLogo check FAILED!")